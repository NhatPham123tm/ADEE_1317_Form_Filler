from docx import Document
import os
# Load the current number from a file
def load_current_number(file_path):
    if os.path.exists(file_path):
        with open(file_path, 'r') as f:
            return int(f.read().strip())
    return 1  # Default starting number

# Save the next number
def save_next_number(file_path, number):
    with open(file_path, 'w') as f:
        f.write(str(number))

def fill_docx(input_docx, output_docx, replacements):
    doc = Document(input_docx)

    for para in doc.paragraphs:
        for key, value in replacements.items():
            if key in para.text:
                para.text = para.text.replace(key, value)

    # Also update in tables (if you have placeholders inside tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)

    doc.save(output_docx)

def digit_control(counter):
    return str(counter).zfill(8)

# Example Usage
template_doc = "ADEE-1317-texas-adult-driver-education-certificate-template.docx"
counter_file = 'counter.txt'

# Load current number
current_number = load_current_number(counter_file)
fill_number = digit_control(current_number)

output_doc = "filled_certificate" + fill_number + ".docx"
                                        

replacements = {
    '{{CONTROL_NUMBER}}': 'DEE ' + fill_number,
}

fill_docx(template_doc, output_doc, replacements)

# Save the next number
save_next_number(counter_file, current_number + 1)

print(f"Created {output_doc} with control number {replacements['{{CONTROL_NUMBER}}']}")
