
from tkinter import filedialog, messagebox
from docx import Document
from docx2pdf import convert
import os

def fill_docx(template_path, output_path, replacements):
    doc = Document(template_path)
    for para in doc.paragraphs:
        for key, value in replacements.items():
            if key in para.text:
                para.text = para.text.replace(key, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)

    doc.save(output_path)

def load_current_number(file_path):
    if os.path.exists(file_path):
        with open(file_path, 'r') as f:
            return int(f.read().strip())
    return 1  # Default starting number

# Save the next number
def save_next_number(file_path, number):
    with open(file_path, 'w') as f:
        f.write(str(number))

def digit_control(counter):
    return str(counter).zfill(8)

def generate_doc(filling):
    counter_file = 'counter.txt'
    current_number = load_current_number(counter_file)
    control_number = digit_control(current_number)  

    template_path = filedialog.askopenfilename(title="Choose a DOCX template", filetypes=[("Word files", "*.docx")])
    if not template_path:
        return

    output_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
    if not output_path:
        return

    replacements = {
        "{{LAST_NAME}}": filling['last_name_entry'].get(),
        "{{FIRST_NAME}}": filling['first_name_entry'].get(),
        "{{ROAD_RULE_POINTS}}": filling['road_rule_entry'].get(),
        "{{ROAD_SIGN_POINTS}}": filling['road_sign_entry'].get(),
        "{{SCHOOL_NAME}}": filling['school_name_entry'].get(),
        "{{TDLR_NUMBER}}": filling['TDLR_entry'].get(),
        "{{EDUCATOR_NUMBER}}": filling['driver_school_number_entry'].get(),
        "{{DATE_ISSUED}}": filling['date_issued_entry'].get(),
        "{{CONTROL_NUMBER}}": "DEE "+ control_number,
    }

    fill_docx(template_path, output_path, replacements)

    save_next_number(counter_file, current_number + 1)

    # Optional: convert to PDF
    if messagebox.askyesno("Convert", "Do you want to export to PDF?"):
        try:
            convert(output_path)  # creates PDF in same folder
        except Exception as e:
            messagebox.showwarning("PDF Error", f"PDF export failed: {e}")

    messagebox.showinfo("Success", f"Filled document saved to:\n{output_path}")
